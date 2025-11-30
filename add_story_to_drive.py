#!/usr/bin/env python3
"""
Add Story to Google Drive and RAG System

APPENDS stories to an existing Google Doc in Google Drive AND uploads to the
Gemini File Search store for RAG indexing.

The Google Doc must be:
1. Created by the user (not the service account)
2. Shared with the service account with Editor access
3. Named "stories" in the "rufftree" folder

This approach works because service accounts can EDIT shared documents but
cannot CREATE new files (zero storage quota for service accounts).

Usage:
    python add_story_to_drive.py --title "Story Title" --about "Patrick Ruff, Jenny Wang" --story "Story content..."
"""

import os
import sys
import json
import argparse
import tempfile
from datetime import datetime
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from google.oauth2 import service_account
from googleapiclient.discovery import build

# Import our existing file search functions
from test_file_search import get_client, get_or_create_store, upload_document


# Configuration
DRIVE_FOLDER_NAME = "rufftree"
STORIES_DOC_NAME = "stories"  # Name of the shared Google Doc to append to
# Need drive scope to find files and docs scope to edit
SCOPES = [
    'https://www.googleapis.com/auth/drive.readonly',
    'https://www.googleapis.com/auth/documents'
]


def get_google_services():
    """Initialize Google Drive and Docs API services."""
    creds_json = os.getenv("GOOGLE_DRIVE_CREDENTIALS")

    if not creds_json:
        raise ValueError(
            "GOOGLE_DRIVE_CREDENTIALS environment variable not set. "
            "This should contain your service account JSON credentials."
        )

    try:
        creds_info = json.loads(creds_json)
        service_account_email = creds_info.get('client_email', 'unknown')
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in GOOGLE_DRIVE_CREDENTIALS: {e}")

    credentials = service_account.Credentials.from_service_account_info(
        creds_info,
        scopes=SCOPES
    )

    drive_service = build('drive', 'v3', credentials=credentials)
    docs_service = build('docs', 'v1', credentials=credentials)

    print("✅ Connected to Google APIs")
    print(f"   Service Account: {service_account_email}")

    return drive_service, docs_service, service_account_email


def find_stories_doc(drive_service, service_account_email: str) -> tuple:
    """
    Find the shared 'stories' Google Doc that the service account can edit.

    Returns: (doc_id, web_link)
    """
    # Search for Google Docs named "stories" that we have access to
    query = f"name='{STORIES_DOC_NAME}' and mimeType='application/vnd.google-apps.document' and trashed=false"

    results = drive_service.files().list(
        q=query,
        spaces='drive',
        fields='files(id, name, webViewLink, owners)',
        supportsAllDrives=True,
        includeItemsFromAllDrives=True
    ).execute()

    files = results.get('files', [])

    if not files:
        print(f"\n❌ ERROR: Google Doc '{STORIES_DOC_NAME}' not found!")
        print(f"\n📋 To fix this, you need to:")
        print(f"   1. Go to your Google Drive (drive.google.com)")
        print(f"   2. Create a Google Doc named '{STORIES_DOC_NAME}'")
        print(f"   3. Move it to your '{DRIVE_FOLDER_NAME}' folder")
        print(f"   4. Right-click the doc → Share")
        print(f"   5. Add this email with 'Editor' access:")
        print(f"      {service_account_email}")
        print(f"   6. Click 'Share'")
        print(f"\n   Then run this workflow again.")
        raise ValueError(f"Google Doc '{STORIES_DOC_NAME}' not found. See instructions above.")

    doc = files[0]
    doc_id = doc['id']
    web_link = doc.get('webViewLink', '')

    print(f"📄 Found stories doc: {STORIES_DOC_NAME} (ID: {doc_id})")
    if web_link:
        print(f"   🔗 Link: {web_link}")

    return doc_id, web_link


def generate_filename(title: str) -> str:
    """Generate filename: yyyymmdd_patstory_title (used for RAG file)"""
    date_str = datetime.now().strftime("%Y%m%d")
    safe_title = "".join(c if c.isalnum() or c in ' -_' else '' for c in title)
    safe_title = safe_title.strip().replace(' ', '_').lower()[:50]
    return f"{date_str}_patstory_{safe_title}"


def append_story_to_doc(docs_service, doc_id: str, title: str,
                        about: str, story: str, author: str) -> None:
    """Append a story to the existing shared Google Doc."""

    # First, get the current document to find where to append
    doc = docs_service.documents().get(documentId=doc_id).execute()

    # Find the end of the document
    content = doc.get('body', {}).get('content', [])
    end_index = 1  # Default to beginning if document is empty

    for element in content:
        if 'endIndex' in element:
            end_index = element['endIndex']

    # Account for the newline at the end
    insert_index = max(1, end_index - 1)

    print(f"   📝 Appending story at index {insert_index}...")

    # Build the story content to append
    date_str = datetime.now().strftime("%B %d, %Y")
    word_count = len(story.split())

    # Create separator and story block
    story_block = f"""

{'═' * 60}

{title}
{'─' * 40}
By: {author}
About: {about}
Date: {date_str}
{'─' * 40}

{story}

Word Count: {word_count}
{'═' * 60}
"""

    # Insert the story block at the end
    requests = [
        {
            'insertText': {
                'location': {'index': insert_index},
                'text': story_block
            }
        }
    ]

    # Execute the batch update
    docs_service.documents().batchUpdate(
        documentId=doc_id,
        body={'requests': requests}
    ).execute()

    print(f"   ✅ Story appended successfully!")


def create_docx_for_rag(title: str, about: str, story: str, author: str) -> BytesIO:
    """Create a Word document for RAG indexing."""
    doc = Document()

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    meta_para.add_run("By: ").bold = True
    meta_para.add_run(author)

    meta_para2 = doc.add_paragraph()
    meta_para2.add_run("About: ").bold = True
    meta_para2.add_run(about)

    meta_para3 = doc.add_paragraph()
    meta_para3.add_run("Date: ").bold = True
    meta_para3.add_run(datetime.now().strftime("%B %d, %Y"))

    doc.add_paragraph("─" * 50)
    doc.add_paragraph()

    story_paragraphs = story.strip().split('\n\n')
    for para_text in story_paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph()
    footer.add_run(f"Word Count: {len(story.split())}").italic = True

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer


def add_story(title: str, about: str, story: str, author: str = "Patrick Ruff",
              decade: str = None, specific_date: str = None):
    """Main function to append story to shared Google Doc AND upload to RAG system."""
    print("📖 Adding Story to Google Drive & RAG System")
    print("=" * 60)
    print(f"Title: {title}")
    print(f"About: {about}")
    print(f"Author: {author}")
    print(f"Word Count: {len(story.split())}")
    print("=" * 60)

    # === Part 1: Append to existing Google Doc in Drive ===
    print("\n📄 STEP 1: Appending to shared Google Doc...")

    drive_service, docs_service, service_account_email = get_google_services()

    # Find the existing "stories" document (must be shared with service account)
    doc_id, web_link = find_stories_doc(drive_service, service_account_email)

    # Append the story to the document
    append_story_to_doc(docs_service, doc_id, title, about, story, author)

    # === Part 2: Upload to File Search Store for RAG ===
    print("\n🔍 STEP 2: Adding to RAG File Search store...")

    docx_buffer = create_docx_for_rag(title, about, story, author)
    rag_filename = generate_filename(title)
    filename = f"{rag_filename}.docx"

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / filename
        with open(temp_path, 'wb') as f:
            f.write(docx_buffer.getvalue())

        print("   Connecting to Gemini File Search...")
        client = get_client()
        store_name = get_or_create_store(client)

        # Build custom metadata for RAG filtering
        # This allows filtering stories by author or subject during queries
        custom_metadata = [
            {"key": "author", "string_value": author},
            {"key": "content_type", "string_value": "story"}
        ]

        # Add each subject as metadata (comma-separated list becomes multiple entries)
        subjects = [s.strip() for s in about.split(',') if s.strip()]
        for subject in subjects:
            custom_metadata.append({"key": "subject_of_story", "string_value": subject})

        # Add decade metadata if provided
        if decade:
            custom_metadata.append({"key": "decade", "string_value": decade})
            # Try to extract a representative year from decade
            import re
            year_match = re.search(r'(\d{4})', decade)
            if year_match:
                year = int(year_match.group(1))
                custom_metadata.append({"key": "year", "numeric_value": year})

        # Add specific date if provided
        if specific_date:
            custom_metadata.append({"key": "specific_date", "string_value": specific_date})

        print(f"   📋 Adding metadata: author={author}, subjects={subjects}"
              + (f", decade={decade}" if decade else ""))
        upload_document(client, store_name, str(temp_path), custom_metadata=custom_metadata)

    # === Summary ===
    print("\n" + "=" * 60)
    print("✅ STORY ADDED SUCCESSFULLY!")
    print("=" * 60)
    print(f"📄 Appended to: {STORIES_DOC_NAME}")
    if web_link:
        print(f"🔗 View: {web_link}")
    print(f"🔍 RAG: Indexed as {filename}")
    print("=" * 60)

    return doc_id, filename


def main():
    parser = argparse.ArgumentParser(description='Add a family story to Google Drive and RAG')
    parser.add_argument('--title', '-t', help='Story title')
    parser.add_argument('--about', '-a', help='Who the story is about (comma-separated names)')
    parser.add_argument('--story', '-s', help='Story content')
    parser.add_argument('--author', default='Patrick Ruff', help='Story author (default: Patrick Ruff)')
    parser.add_argument('--decade', help='Decade when story took place (e.g., "1980s (1980-1989)")')
    parser.add_argument('--specific-date', dest='specific_date', help='Specific date/year (e.g., "Summer 1985")')
    parser.add_argument('--interactive', '-i', action='store_true', help='Interactive mode')

    args = parser.parse_args()

    decade = None
    specific_date = None

    if args.interactive:
        print("📖 Add Family Story")
        print("=" * 60)
        title = input("\nStory Title: ").strip()
        about = input("Who is this story about? (comma-separated): ").strip()
        author = input("Your name (press Enter for 'Patrick Ruff'): ").strip() or "Patrick Ruff"
        decade = input("Decade (e.g., '1980s', press Enter to skip): ").strip() or None
        specific_date = input("Specific date/year (e.g., 'Summer 1985', press Enter to skip): ").strip() or None
        print("\nEnter your story (end with two empty lines):")
        lines = []
        empty_count = 0
        while True:
            try:
                line = input()
                if line == "":
                    empty_count += 1
                    if empty_count >= 2:
                        break
                    lines.append(line)
                else:
                    empty_count = 0
                    lines.append(line)
            except EOFError:
                break
        story = "\n".join(lines).strip()
    else:
        if not args.title or not args.about or not args.story:
            parser.error("--title, --about, and --story are required (or use --interactive)")

        title = args.title
        about = args.about
        story = args.story
        author = args.author
        decade = args.decade
        specific_date = args.specific_date

    if not title or not about or not story:
        print("❌ Error: Title, about, and story content are required")
        sys.exit(1)

    try:
        add_story(title, about, story, author, decade=decade, specific_date=specific_date)
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
